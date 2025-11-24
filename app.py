import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import re
import datetime

# ==========================================
# 1. KONFIGURASI HALAMAN & CSS
# ==========================================
st.set_page_config(
    page_title="Generator RPP - Baskoro Pandu Aji",
    layout="wide",
    initial_sidebar_state="expanded"
)

# CSS Custom (Classic Green - CLEAN TANPA IKON)
st.markdown("""
<style>
    .stApp { background-color: #f1f8f3; }
    h1 { color: #1b5e20 !important; font-family: 'Helvetica', sans-serif; }
    h2, h3 { color: #2e7d32 !important; }
    
    div[data-testid="stForm"] {
        background-color: #ffffff;
        padding: 30px;
        border-radius: 15px;
        box-shadow: 0 4px 6px rgba(0,0,0,0.05);
        border-top: 5px solid #4caf50;
    }
    
    div.stButton > button {
        background-color: #2e7d32;
        color: white;
        border-radius: 8px;
        padding: 10px 24px;
        font-weight: bold; 
        border: none;
        width: 100%;
        text-transform: uppercase;
    }
    div.stButton > button:hover { background-color: #1b5e20; }
    
    /* Profil Sidebar */
    .profile-img {
        display: block; margin-left: auto; margin-right: auto;
        width: 140px; height: 140px; object-fit: cover;
        border-radius: 50%; border: 4px solid #4caf50; margin-bottom: 15px;
    }
    .profile-name {
        text-align: center; font-weight: bold; font-size: 16px; color: #1b5e20; margin-bottom: 5px;
    }
    .profile-desc {
        text-align: center; font-size: 12px; color: #555; margin-bottom: 20px;
    }
    
    /* Tombol Link Sidebar (Clean Text) */
    .clean-btn {
        display: block; text-align: center; background-color: white;
        color: #2e7d32; padding: 10px; margin: 8px 0; border-radius: 6px;
        text-decoration: none; font-weight: 600; border: 1px solid #c8e6c9;
        font-size: 14px; transition: 0.3s;
    }
    .clean-btn:hover { background-color: #e8f5e9; color: #1b5e20; border-color: #2e7d32; }

    /* Tombol Utama Profil (Solid Green) */
    .main-profile-btn {
        display: block; text-align: center; background-color: #1b5e20;
        color: white !important; padding: 12px; margin: 15px 0; border-radius: 6px;
        text-decoration: none; font-weight: bold; border: none;
        font-size: 14px; letter-spacing: 0.5px;
        transition: 0.3s;
    }
    .main-profile-btn:hover { background-color: #0F4D2A; opacity: 0.9; }

    /* Kotak Donasi (Clean) */
    .donation-box {
        background-color: #fff8e1; border: 1px solid #ffeeba;
        color: #856404; padding: 15px; border-radius: 8px;
        text-align: center; margin-top: 20px; font-size: 13px;
    }
    .donation-number {
        font-size: 18px; font-weight: bold; color: #1b5e20; margin: 5px 0;
    }
    
    .footer-text { text-align: center; color: #888; font-size: 11px; margin-top: 40px; }
</style>
""", unsafe_allow_html=True)

# ==========================================
# 2. API SETUP
# ==========================================
# Mengambil API Key dari Secrets (Brankas Server Streamlit)
try:
    api_key = st.secrets["GEMINI_API_KEY"]
except:
    # Cadangan kosong agar aman di GitHub
    api_key = "" 

if api_key:
    genai.configure(api_key=api_key)

# ==========================================
# 3. SIDEBAR (CLEAN TEXT ONLY)
# ==========================================
with st.sidebar:
    st.markdown("""
        <img src="https://i.ibb.co.com/00QNsb6/IMG-6622-2.jpg" class="profile-img">
        <div class="profile-name">Baskoro Pandu Aji, S.Pd.</div>
    """, unsafe_allow_html=True)
    
    st.markdown("""
    <a href="https://baskoropandu.my.canva.site/profil-baskoro-pandu-aji" target="_blank" class="main-profile-btn">
        KUNJUNGI PROFIL
    </a>
    """, unsafe_allow_html=True)

    st.markdown("""
    <div class="profile-desc">
        Pengembang Aplikasi Pendidikan &<br>Kepala Sekolah SDN 183/II Sumber Mulya
    </div>
    """, unsafe_allow_html=True)
    
    st.write("---")
    
    st.markdown("**KONTAK SAYA**")
    st.markdown("""
    <a href="https://wa.me/6281329951000" target="_blank" class="clean-btn">WhatsApp</a>
    <a href="https://instagram.com/baskoro_panduaji" target="_blank" class="clean-btn">Instagram</a>
    """, unsafe_allow_html=True)
    
    st.markdown("""
    <div class="donation-box">
        DUKUNG PENGEMBANGAN<br>
        Top Up GoPay:<br>
        <div class="donation-number">081329951000</div>
        a.n Baskoro Pandu Aji
    </div>
    """, unsafe_allow_html=True)

    st.write("---")
    if st.button("RESET DATA"):
        st.session_state.clear()
        st.rerun()

# ==========================================
# 4. LOGIKA AI (RPM DEEP LEARNING STRUKTUR)
# ==========================================
def get_best_model():
    try:
        target_models = ['gemini-1.5-flash', 'gemini-1.5-pro', 'gemini-pro']
        available_models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
        for target in target_models:
            for available in available_models:
                if target in available: return available
        return available_models[0] if available_models else None
    except: return "gemini-pro"

def generate_rpp_deep_learning(data):
    model_name = get_best_model()
    if not model_name: return "Error: Masalah koneksi API."
    
    profil_str = ", ".join(data['profil'])
    asesmen_str = ", ".join(data['asesmen'])
    
    prompt = f"""
    Bertindaklah sebagai **Guru Profesional**. Buat **Rancangan Pembelajaran Mendalam (RPM)**.
    
    INSTRUKSI:
    1.  Langsung ke konten RPP (Tanpa kalimat pembuka).
    2.  Identitas TIDAK PERLU (sudah ada tabel otomatis).
    3.  **Wajib** sertakan: CP, Mitra, Lingkungan, Digital.
    4.  **Kegiatan Pembelajaran:** Format LIST (Poin-poin) mengandung Sintaks {data['metode']}.
    
    DATA:
    Mapel: {data['mapel']} | Topik: {data['topik']} | TP: {data['tp']} | Model: {data['metode']}
    
    STRUKTUR OUTPUT MARKDOWN:
    
    ## A. CAPAIAN PEMBELAJARAN (CP)
    (Tuliskan CP yang relevan secara singkat)

    ## B. KOMPONEN PERANCANGAN PEMBELAJARAN
    **1. Mitra Pembelajaran**
    (Pihak yang dilibatkan: Orang tua, Pakar, dll)
    **2. Lingkungan Pembelajaran**
    (Setting kelas/luar kelas)
    **3. Pemanfaatan Digital**
    (Aplikasi/Platform yang digunakan)
    
    ## C. KOMPONEN INTI
    **1. Tujuan Pembelajaran**
    {data['tp']}
    **2. Pemahaman Bermakna**
    (Manfaat kontekstual)
    **3. Pertanyaan Pemantik**
    (Pertanyaan HOTS)
    **4. Materi Esensial**
    (Poin materi kunci)
    
    ## D. KEGIATAN PEMBELAJARAN
    ### 1. Pendahuluan (10 Menit)
    * (Sapaan, Doa, Presensi, Apersepsi, Tujuan)
    
    ### 2. Kegiatan Inti (... Menit) - Model: {data['metode']}
    (Uraikan langkah demi langkah berdasarkan SINTAKS {data['metode']}. Fokus pada aktivitas siswa aktif).
    
    ### 3. Penutup (10 Menit)
    * (Refleksi, Kesimpulan, Doa)
    
    ## E. ASESMEN
    * **Rencana Asesmen:** {asesmen_str}
    * **Formatif:** (Instrumen/Teknik)
    * **Sumatif:** (Instrumen/Teknik)
    
    ## F. MEDIA & SUMBER BELAJAR
    * **Sumber Belajar Utama:** Buku Paket.
    * **Sumber Belajar Digital (Pengayaan):**
      - **Video:** (Berikan Judul Video Youtube yang Relevan)
      - **Bacaan:** (Berikan Judul Artikel Web)
    
    ## G. LAMPIRAN
    ### 1. Lembar Kerja Peserta Didik (LKPD)
    (Rancangan tugas)
    ### 2. Rubrik Penilaian
    (Gunakan Tabel Markdown)
    | Kriteria | Skor 4 (Sangat Baik) | Skor 3 (Baik) | Skor 2 (Cukup) | Skor 1 (Kurang) |
    |---|---|---|---|---|
    | (Isi Kriteria) | ... | ... | ... | ... |
    """
    try:
        model = genai.GenerativeModel(model_name)
        response = model.generate_content(prompt)
        text = response.text
        
        # Bersihkan basa-basi & format HTML
        text = re.sub(r'^(Tentu|Baik|Berikut|Sesuai).*?(\n\n|\n)', '', text, flags=re.IGNORECASE).strip()
        text = text.replace("<br>", "\n").replace("<br/>", "\n")
        return text
    except Exception as e: return f"Error: {e}"

# ==========================================
# 5. FORMATTER WORD (TABEL IDENTITAS + SIGNATURE)
# ==========================================
def markdown_to_word_table(doc, table_text):
    lines = table_text.strip().split('\n')
    lines = [l for l in lines if not set(l.strip().replace('|','').replace('-','').replace(' ','')) == set()]
    if len(lines) < 2: return 
    headers = [c.strip() for c in lines[0].strip('|').split('|')]
    col_count = len(headers)
    table = doc.add_table(rows=0, cols=col_count)
    table.style = 'Table Grid'
    hdr_row = table.add_row().cells
    for i, h in enumerate(headers):
        hdr_row[i].text = h
        for p in hdr_row[i].paragraphs:
            for r in p.runs: r.bold = True
    for line in lines[1:]:
        if "|" not in line: continue
        row_cells = table.add_row().cells
        cells = [c.strip() for c in line.strip('|').split('|')]
        for i, c in enumerate(cells):
            if i < col_count: row_cells[i].text = c

def create_docx_formatted(rpp_text, data):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # HEADER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"RANCANGAN PEMBELAJARAN MENDALAM")
    r.bold = True; r.font.size = Pt(14)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"MATA PELAJARAN: {data['mapel'].upper()}")
    r2.bold = True
    doc.add_paragraph()

    # IDENTITAS (TABEL)
    doc.add_heading('INFORMASI UMUM', level=1).style.font.color.rgb = RGBColor(0,0,0)
    table = doc.add_table(rows=8, cols=3)
    table.style = 'Table Grid'
    for row in table.rows:
        row.cells[0].width = Inches(1.8); row.cells[1].width = Inches(0.2); row.cells[2].width = Inches(4.5)
    info = [
        ("Nama Penyusun", ":", data['guru']), ("Satuan Pendidikan", ":", data['sekolah']),
        ("Kelas / Fase", ":", f"{data['kelas']} / {data['fase']}"), ("Tahun Pelajaran", ":", f"{data['tahun']} ({data['semester']})"),
        ("Topik / Materi", ":", data['topik']), ("Alokasi Waktu", ":", data['waktu']),
        ("Model Pembelajaran", ":", data['metode']), ("Profil Lulusan", ":", ", ".join(data['profil']))
    ]
    for i, item in enumerate(info):
        cells = table.rows[i].cells
        cells[0].text = item[0]; cells[0].paragraphs[0].runs[0].bold = True
        cells[1].text = item[1]; cells[2].text = item[2]
    doc.add_paragraph()

    # PARSING ISI
    lines = rpp_text.split('\n')
    table_buffer = []
    in_table = False
    for line in lines:
        line = line.strip()
        if not line: continue
        if line.startswith("|"):
            in_table = True; table_buffer.append(line); continue
        else:
            if in_table:
                markdown_to_word_table(doc, "\n".join(table_buffer))
                doc.add_paragraph(); table_buffer = []; in_table = False
        
        if line.startswith('# '): doc.add_heading(line.replace('# ', ''), level=1).style.font.color.rgb = RGBColor(0,0,0)
        elif line.startswith('## '): doc.add_heading(line.replace('## ', ''), level=2).style.font.color.rgb = RGBColor(0,0,0)
        elif line.startswith('### '): doc.add_heading(line.replace('### ', ''), level=3).style.font.color.rgb = RGBColor(0,0,0)
        elif line.startswith('* ') or line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            clean = line[2:]; parts = re.split(r'(\*\*.*?\*\*)', clean)
            for part in parts:
                if part.startswith('**') and part.endswith('**'): p.add_run(part.replace('**', '')).bold = True
                else: p.add_run(part)
        else:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'): p.add_run(part.replace('**', '')).bold = True
                else: p.add_run(part)
    if in_table and table_buffer: markdown_to_word_table(doc, "\n".join(table_buffer))

    # --- PENGESAHAN (TANDA TANGAN) ---
    doc.add_paragraph()
    doc.add_paragraph() # Spasi
    
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.autofit = True
    
    # Kiri (Kepsek)
    left_cell = sig_table.rows[0].cells[0]
    p = left_cell.add_paragraph("Mengetahui,\nKepala Sekolah")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = left_cell.add_paragraph("\n\n\n")
    p = left_cell.add_paragraph(data['kepsek'])
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True; p.runs[0].underline = True
    p = left_cell.add_paragraph(f"NIP. {data['nip_kepsek']}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Kanan (Guru)
    right_cell = sig_table.rows[0].cells[1]
    p = right_cell.add_paragraph(f"{data['kota']}, {data['tanggal']}\nGuru Mata Pelajaran")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = right_cell.add_paragraph("\n\n\n")
    p = right_cell.add_paragraph(data['guru'])
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True; p.runs[0].underline = True
    p = right_cell.add_paragraph(f"NIP. {data['nip_guru']}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    bio = io.BytesIO()
    doc.save(bio)
    return bio

# ==========================================
# 6. UI UTAMA
# ==========================================
st.title("GENERATOR RPP PINTAR")
st.markdown("**Kurikulum Merdeka & Pembelajaran Mendalam**")

with st.form("form_rpp"):
    st.markdown("### 1. IDENTITAS & PENGESAHAN")
    c1, c2 = st.columns(2)
    with c1:
        guru = st.text_input("Nama Guru", value="Tuti Haryanti, S.Pd.")
        nip_guru = st.text_input("NIP Guru", placeholder="19xxxxxxxxxxxx")
        sekolah = st.text_input("Nama Sekolah", placeholder="Contoh: SDN 183/II Sumber Mulya")
        kota = st.text_input("Kota/Kabupaten", value="Jambi")
        
    with c2:
        kepsek = st.text_input("Nama Kepala Sekolah", value="Baskoro Pandu Aji, S.Pd.")
        nip_kepsek = st.text_input("NIP Kepala Sekolah", placeholder="19xxxxxxxxxxxx")
        tanggal = st.text_input("Tanggal RPP", value=datetime.date.today().strftime("%d %B %Y"))
        
    st.markdown("---")
    st.markdown("### 2. DATA PEMBELAJARAN")
    c3, c4 = st.columns(2)
    with c3:
        mapel = st.text_input("Mata Pelajaran", placeholder="Koding dan Kecerdasan Artifisial")
        kelas = st.selectbox("Kelas", [str(i) for i in range(1, 13)])
        fase = st.selectbox("Fase", ["A (SD 1-2)", "B (SD 3-4)", "C (SD 5-6)", "D (SMP)", "E (SMA 10)", "F (SMA 11-12)"])
        tahun = st.text_input("Tahun Pelajaran", value="2024/2025")
        semester = st.selectbox("Semester", ["Ganjil", "Genap"])
        
    with c4:
        topik = st.text_input("Topik / Materi Pembelajaran")
        waktu = st.text_input("Alokasi Waktu", value="2 x 35 menit")
        metode = st.selectbox("Model Pembelajaran", ["Problem Based Learning (PBL)", "Project Based Learning (PjBL)", "Inquiry Learning", "Discovery Learning", "Cooperative Learning"])
        
    tp = st.text_area("Tujuan Pembelajaran (TP)", height=80)
    
    st.markdown("### 3. ASESMEN & PROFIL")
    c5, c6 = st.columns(2)
    with c5:
        asesmen = st.multiselect("Bentuk Asesmen", ["Tes Tertulis", "Tes Lisan", "Penugasan/Proyek", "Portofolio", "Kinerja/Praktik", "Observasi"], default=["Tes Tertulis", "Penugasan/Proyek"])
    with c6:
        profil = st.multiselect("Dimensi Profil Lulusan", ["Keimanan & Ketakwaan thd Tuhan YME", "Kewargaan", "Penalaran Kritis", "Kreativitas", "Kolaborasi", "Kemandirian", "Kesehatan", "Komunikasi"], default=["Penalaran Kritis", "Kreativitas"])

    submitted = st.form_submit_button("BUAT RPP SEKARANG")

if submitted:
    if not api_key: st.error("⚠️ API Key belum terdeteksi!")
    elif not topik: st.warning("⚠️ Mohon isi Topik materi terlebih dahulu.")
    else:
        with st.spinner('Sedang menyusun RPP Lengkap...'):
            data = {
                "guru":guru, "nip_guru":nip_guru, "kepsek":kepsek, "nip_kepsek":nip_kepsek,
                "sekolah":sekolah, "kota":kota, "tanggal":tanggal,
                "mapel":mapel, "kelas":kelas, "fase":fase, "waktu":waktu,
                "tahun":tahun, "semester":semester, "topik":topik, "tp":tp,
                "metode":metode, "profil":profil, "asesmen":asesmen
            }
            
            hasil = generate_rpp_deep_learning(data)
            if "Error" in hasil:
                st.error(hasil)
            else:
                st.success("✅ RPP Berhasil Dibuat!")
                with st.expander("📄 Lihat Hasil (Preview)", expanded=True): st.markdown(hasil)
                docx = create_docx_formatted(hasil, data)
                st.download_button("📥 UNDUH FILE WORD (.DOCX)", docx.getvalue(), f"RPP_{mapel}_{topik}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

st.markdown("""<div class="footer-text"><hr><p>Copyright © 2025 <b>Baskoro Pandu Aji, S.Pd.</b><br>All Rights Reserved.</p></div>""", unsafe_allow_html=True)
